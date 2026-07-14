"""输出格式模块 — Markdown / JSON / Word / HTML"""
import os
import json
from datetime import datetime


def save_markdown(results: list, output_dir: str, prefix: str = "analysis") -> str:
    """保存为 Markdown"""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    path = os.path.join(output_dir, f"{prefix}_{ts}.md")
    with open(path, "w", encoding="utf-8") as f:
        f.write(f"# 案例精华摘录\n\n")
        f.write(f"_生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}，")
        f.write(f"共 {len(results)} 篇_\n\n")
        for r in results:
            t = r.get("title", "未知案例")
            f.write(f"## {t}\n\n")
            f.write(f"**研究背景**\n\n{r.get('background','未提取到')}\n\n")
            f.write(f"**研究对象/目的**\n\n{r.get('objective','未提取到')}\n\n")
            f.write(f"**研究方法**\n\n{r.get('methods','未提取到')}\n\n")
            f.write(f"**实验/实证**\n\n{r.get('experiment','未提取到')}\n\n")
            f.write(f"**核心发现/成果**\n\n{r.get('findings','未提取到')}\n\n")
            f.write(f"**创新点**\n\n{r.get('innovation','未提取到')}\n\n")
            f.write(f"**优势与局限**\n\n{r.get('pros_cons','未提取到')}\n\n")
            f.write(f"**一句话总结**\n\n{r.get('summary','未提取到')}\n\n")
            f.write("---\n\n")
    return path


def save_json(results: list, output_dir: str, prefix: str = "analysis") -> str:
    """保存为 JSON"""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    path = os.path.join(output_dir, f"{prefix}_{ts}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)
    return path


def _set_cjk_font(run, name: str = "Microsoft YaHei", size=None):
    """设置中文字体（解决 Word 乱码问题）"""
    run.font.name = name
    # 设置东亚字体（Word 用这个渲染中文）
    from docx.oxml.ns import qn
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from lxml import etree
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = size


def save_docx(results: list, output_dir: str, prefix: str = "analysis") -> str:
    """保存为 Word 文档 (.docx)

    需要: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        print("⚠  需要安装 python-docx: pip install python-docx")
        return ""

    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    path = os.path.join(output_dir, f"{prefix}_{ts}.docx")

    doc = Document()

    # ── 设置文档默认字体 ──
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # 标题
    title = doc.add_heading('案例精华摘录', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        _set_cjk_font(run, "Microsoft YaHei")

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}　共 {len(results)} 篇"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)
    _set_cjk_font(meta_run, "Microsoft YaHei", Pt(10))

    for r in results:
        h1 = doc.add_heading(r.get("title", "未知案例"), level=1)
        for run in h1.runs:
            _set_cjk_font(run, "Microsoft YaHei")

        for section_name, key in [
            ("研究背景", "background"),
            ("研究对象/目的", "objective"),
            ("研究方法", "methods"),
            ("实验/实证", "experiment"),
            ("核心发现/成果", "findings"),
            ("创新点", "innovation"),
            ("优势与局限", "pros_cons"),
            ("一句话总结", "summary"),
        ]:
            h2 = doc.add_heading(section_name, level=2)
            for run in h2.runs:
                _set_cjk_font(run, "Microsoft YaHei")

            p = doc.add_paragraph()
            prun = p.add_run(r.get(key, "未提取到"))
            _set_cjk_font(prun, "Microsoft YaHei", Pt(11))

        doc.add_paragraph("—" * 30)

    doc.save(path)
    return path


def save_html(results: list, output_dir: str, prefix: str = "analysis") -> str:
    """保存为 HTML 网页，带排版样式"""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    path = os.path.join(output_dir, f"{prefix}_{ts}.html")

    sections = []
    for r in results:
        sections.append(f"""
    <article>
      <h2>{r.get('title', '未知案例')}</h2>
      <div class="meta">模型：{r.get('model', '')} | 方法：{r.get('method', '')}</div>
      <h3>🌍 研究背景</h3>
      <p>{r.get('background', '未提取到')}</p>
      <h3>🎯 研究对象/目的</h3>
      <p>{r.get('objective', '未提取到')}</p>
      <h3>🧪 研究方法</h3>
      <p>{r.get('methods', '未提取到')}</p>
      <h3>🔬 实验/实证</h3>
      <p>{r.get('experiment', '未提取到')}</p>
      <h3>📈 核心发现/成果</h3>
      <p>{r.get('findings', '未提取到')}</p>
      <h3>💡 创新点</h3>
      <p>{r.get('innovation', '未提取到')}</p>
      <h3>⚖️ 优势与局限</h3>
      <p>{r.get('pros_cons', '未提取到')}</p>
      <h3>✨ 一句话总结</h3>
      <p class="summary">{r.get('summary', '未提取到')}</p>
    </article>""")

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>CaseForge · 论文拆解报告</title>
<style>
  :root {{
    --bg: #fafaf9; --card: #fff; --ink: #1c1c1c; --muted: #6b6b6b;
    --accent: #5c6b4e; --border: #e5e1da; --tag-bg: #ece8e0;
  }}
  body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif;
         max-width: 820px; margin: 0 auto; padding: 40px 24px;
         background: var(--bg); color: var(--ink); line-height: 1.8; }}
  /* Header */
  .brand {{ text-align: center; margin-bottom: 32px; }}
  .brand h1 {{ font-size: 32px; font-weight: 800; letter-spacing: -0.02em; margin: 0; }}
  .brand .tagline {{ font-size: 13px; color: var(--muted); margin-top: 4px; }}
  /* Badges */
  .badges {{ display: flex; gap: 8px; flex-wrap: wrap; justify-content: center; margin-bottom: 36px; }}
  .badge {{ font-size: 10px; font-weight: 600; letter-spacing: 0.08em; text-transform: uppercase;
            padding: 4px 10px; border-radius: 3px; background: var(--tag-bg); color: var(--muted); }}
  .badge.ai {{ background: #e8edd8; color: #4a5e24; }}
  /* Report card */
  article {{ background: var(--card); border: 1px solid var(--border); border-radius: 8px;
             padding: 40px; margin-bottom: 24px; }}
  article h2 {{ font-size: 24px; font-weight: 700; margin: 0 0 4px 0; letter-spacing: -0.01em; }}
  .meta {{ font-size: 12px; color: #aaa; margin-bottom: 28px; }}
  article h3 {{ font-size: 14px; font-weight: 700; margin: 28px 0 8px 0; color: #333;
                border-left: 3px solid var(--accent); padding-left: 12px; }}
  article p {{ font-size: 14px; color: #555; margin: 0 0 16px 0; line-height: 1.85; }}
  .summary {{ background: #f5f2ec; padding: 14px 18px; border-radius: 6px; font-weight: 500;
              border-left: 3px solid var(--accent); margin-top: 28px; }}
  /* Footer */
  footer {{ text-align: center; padding: 24px 0; font-size: 11px; color: #bbb; }}
  footer a {{ color: #999; text-decoration: none; }}
  footer a:hover {{ color: var(--accent); }}
</style>
</head>
<body>
<div class="brand">
  <h1>CaseForge</h1>
  <p class="tagline">Turn papers into structured knowledge.</p>
</div>
<div class="badges">
  <span class="badge ai">{results[0].get('model', 'AI') if results else 'AI'}</span>
  <span class="badge">AI Extraction</span>
  <span class="badge">HTML Report</span>
  <span class="badge">Markdown Ready</span>
  <span class="badge">Generated {datetime.now().strftime('%Y-%m-%d')}</span>
</div>
{''.join(sections)}
<footer>
  <p>Generated by <a href="https://github.com/nowornever17/CaseForge">CaseForge</a> · v2.0</p>
</footer>
</body>
</html>"""

    with open(path, "w", encoding="utf-8") as f:
        f.write(html)
    return path


def save_all(results: list, output_dir: str, prefix: str = "analysis") -> dict:
    """保存为所有支持格式"""
    paths = {
        "html": save_html(results, output_dir, prefix),
        "md": save_markdown(results, output_dir, prefix),
        "json": save_json(results, output_dir, prefix),
    }
    docx_path = save_docx(results, output_dir, prefix)
    if docx_path:
        paths["docx"] = docx_path

    print(f"\n  💾  结果已保存：")
    for fmt, p in paths.items():
        print(f"      {fmt.upper():6s} → {p}")
    return paths
